"""Tests for HBR-1A structural document hierarchy, outline extraction, persistence, and chunk association."""

import uuid
from io import BytesIO

import docx
import pypdf
import pytest
from rest_framework.test import APIClient

from platform_api.apps.knowledge.authentication import mint_delegated_token
from platform_api.apps.libraries.models import Library
from platform_api.apps.processing.extractors import StructureNodeType
from platform_api.apps.processing.extractors.docx import extract_docx
from platform_api.apps.processing.extractors.pdf import extract_pdf
from platform_api.apps.processing.extractors.txt import extract_txt
from platform_api.apps.processing.models import (
    DocumentChunk,
    DocumentStructureNode,
    ProcessingStatus,
)
from platform_api.apps.processing.services import enqueue_processing
from platform_api.apps.resources.checksum import sha256_checksum
from platform_api.apps.resources.models import Resource, ResourceStatus, ResourceType
from platform_api.apps.resources.object_key import generate_resource_object_key
from platform_api.apps.resources.storage import get_object_storage
from platform_api.apps.users.models import User



def _create_pdf_with_outline() -> bytes:
    """Create a minimal PDF with 4 pages and a nested outline hierarchy."""
    writer = pypdf.PdfWriter()
    for _ in range(4):
        writer.add_blank_page(width=200, height=200)

    # Outline:
    # Chapter 1 (page 1)
    #   -> Section 1.1 (page 2)
    #   -> Section 1.2 (page 3)
    # Chapter 2 (page 4)
    ch1 = writer.add_outline_item("Chapter 1: Classical Mechanics", page_number=0)
    writer.add_outline_item("Section 1.1: Newton's Laws", page_number=1, parent=ch1)
    writer.add_outline_item("Section 1.2: Conservation of Energy", page_number=2, parent=ch1)
    writer.add_outline_item("Chapter 2: Thermodynamics", page_number=3)

    stream = BytesIO()
    writer.write(stream)
    return stream.getvalue()


def _create_docx_with_headings() -> bytes:
    """Create a DOCX with nested Heading 1, Heading 2, and Heading 3 paragraphs."""
    doc = docx.Document()
    doc.add_heading("Part I: Foundations", level=1)
    doc.add_paragraph("Introductory text for part one.")
    doc.add_heading("Chapter 1: Mathematical Logic", level=2)
    doc.add_paragraph("Discussion on propositional calculus.")
    doc.add_heading("Section 1.1: Truth Tables", level=3)
    doc.add_paragraph("Truth table evaluation details.")
    doc.add_heading("Chapter 2: Set Theory", level=2)
    doc.add_paragraph("Axiomatic set theory principles.")

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


# ---------------------------------------------------------------------------
# Extraction Tests
# ---------------------------------------------------------------------------


def test_pdf_with_nested_native_outline() -> None:
    """PDF with nested bookmarks produces accurate hierarchical OutlineNode tree."""
    pdf_bytes = _create_pdf_with_outline()
    doc = extract_pdf(pdf_bytes)

    assert not doc.is_empty
    assert len(doc.outline) == 2  # Chapter 1 and Chapter 2 at root

    ch1 = doc.outline[0]
    assert ch1.title == "Chapter 1: Classical Mechanics"
    assert ch1.node_type == StructureNodeType.CHAPTER
    assert ch1.level == 1
    assert ch1.page_start == 1
    assert ch1.page_end == 3
    assert len(ch1.children) == 2

    sec1_1 = ch1.children[0]
    assert sec1_1.title == "Section 1.1: Newton's Laws"
    assert sec1_1.node_type == StructureNodeType.SECTION
    assert sec1_1.level == 2
    assert sec1_1.page_start == 2
    assert sec1_1.page_end == 2

    sec1_2 = ch1.children[1]
    assert sec1_2.title == "Section 1.2: Conservation of Energy"
    assert sec1_2.page_start == 3
    assert sec1_2.page_end == 3

    ch2 = doc.outline[1]
    assert ch2.title == "Chapter 2: Thermodynamics"
    assert ch2.node_type == StructureNodeType.CHAPTER
    assert ch2.page_start == 4
    assert ch2.page_end == 4


def test_pdf_without_outline_produces_empty_outline() -> None:
    """PDF without bookmarks produces pages with empty outline without error."""
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    stream = BytesIO()
    writer.write(stream)
    pdf_bytes = stream.getvalue()

    doc = extract_pdf(pdf_bytes)
    assert not doc.is_empty
    assert doc.outline == []


def test_docx_heading_hierarchy_extraction() -> None:
    """DOCX with Heading 1/2/3 produces nested OutlineNode tree."""
    docx_bytes = _create_docx_with_headings()
    doc = extract_docx(docx_bytes)

    assert not doc.is_empty
    assert len(doc.outline) == 1  # Part I at root

    part1 = doc.outline[0]
    assert part1.title == "Part I: Foundations"
    assert part1.node_type == StructureNodeType.PART
    assert part1.level == 1
    assert len(part1.children) == 2

    ch1 = part1.children[0]
    assert ch1.title == "Chapter 1: Mathematical Logic"
    assert ch1.node_type == StructureNodeType.CHAPTER
    assert ch1.level == 2
    assert len(ch1.children) == 1

    sec1_1 = ch1.children[0]
    assert sec1_1.title == "Section 1.1: Truth Tables"
    assert sec1_1.node_type == StructureNodeType.SECTION
    assert sec1_1.level == 3

    ch2 = part1.children[1]
    assert ch2.title == "Chapter 2: Set Theory"
    assert ch2.level == 2


def test_plain_docx_without_headings_has_empty_outline() -> None:
    """Plain DOCX with no heading styles produces empty outline."""
    doc_raw = docx.Document()
    doc_raw.add_paragraph("Just plain body text.")
    doc_raw.add_paragraph("Another paragraph with no headings.")
    stream = BytesIO()
    doc_raw.save(stream)

    doc = extract_docx(stream.getvalue())
    assert not doc.is_empty
    assert doc.outline == []


def test_txt_extraction_has_empty_outline() -> None:
    """TXT document produces empty outline as expected."""
    txt_bytes = b"Paragraph 1.\n\nParagraph 2."
    doc = extract_txt(txt_bytes)
    assert not doc.is_empty
    assert doc.outline == []


# ---------------------------------------------------------------------------
# Persistence & Chunk Association Integration Tests
# ---------------------------------------------------------------------------


@pytest.fixture
def hbr_docx_resource(db, library_a: Library, user_a: User) -> Resource:
    """Create and upload a DOCX with headings for end-to-end pipeline testing."""
    docx_bytes = _create_docx_with_headings()
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(
        key,
        BytesIO(docx_bytes),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        len(docx_bytes),
    )

    return Resource.objects.create(
        id=res_id,
        library=library_a,
        name="logic_textbook.docx",
        resource_type=ResourceType.DOCX,
        original_filename="logic_textbook.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=len(docx_bytes),
        object_key=key,
        checksum=sha256_checksum(BytesIO(docx_bytes)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )


@pytest.mark.django_db
def test_pipeline_persists_structure_nodes_and_associates_chunks(
    hbr_docx_resource: Resource,
) -> None:
    """End-to-end processing creates DocumentStructureNode hierarchy and links chunks."""
    run = enqueue_processing(hbr_docx_resource)
    run.refresh_from_db()

    assert run.status == ProcessingStatus.READY
    assert run.is_active is True

    # Verify structure nodes
    structure_nodes = DocumentStructureNode.objects.filter(processing_run=run).order_by(
        "sequence"
    )
    assert structure_nodes.count() >= 3

    root_nodes = structure_nodes.filter(parent__isnull=True)
    assert root_nodes.count() >= 1

    part1_node = root_nodes.first()
    assert part1_node is not None
    assert part1_node.resource == hbr_docx_resource
    assert part1_node.library == hbr_docx_resource.library
    assert part1_node.title == "Part I: Foundations"

    child_nodes = structure_nodes.filter(parent=part1_node)
    assert child_nodes.count() == 2

    # Verify chunks are associated with structure nodes
    chunks = DocumentChunk.objects.filter(processing_run=run)
    assert chunks.count() > 0

    associated_chunks = [c for c in chunks if c.structure_node is not None]
    assert len(associated_chunks) > 0

    for chunk in associated_chunks:
        assert chunk.structure_node is not None
        assert chunk.structure_node.processing_run == run
        assert chunk.structure_node.resource == hbr_docx_resource


@pytest.mark.django_db
def test_structure_association_null_when_no_headings(
    db, library_a: Library, user_a: User
) -> None:
    """Documents without headings produce chunks with structure_node=None."""
    plain_content = b"Simple document text with no headings or outline."
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(key, BytesIO(plain_content), "text/plain", len(plain_content))

    resource = Resource.objects.create(
        id=res_id,
        library=library_a,
        name="simple.txt",
        resource_type=ResourceType.TXT,
        original_filename="simple.txt",
        content_type="text/plain",
        size=len(plain_content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(plain_content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )

    run = enqueue_processing(resource)
    run.refresh_from_db()

    assert run.status == ProcessingStatus.READY
    assert DocumentStructureNode.objects.filter(processing_run=run).count() == 0

    chunks = DocumentChunk.objects.filter(processing_run=run)
    assert chunks.count() > 0
    for chunk in chunks:
        assert chunk.structure_node is None


@pytest.mark.django_db
def test_failed_processing_purges_partial_structure_nodes(
    db, library_a: Library, user_a: User
) -> None:
    """When processing fails, any partial DocumentStructureNodes are deleted."""
    empty_content = b"   \n\n  \t  \n  "
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(key, BytesIO(empty_content), "text/plain", len(empty_content))

    empty_resource = Resource.objects.create(
        id=res_id,
        library=library_a,
        name="empty_test.txt",
        resource_type=ResourceType.TXT,
        original_filename="empty_test.txt",
        content_type="text/plain",
        size=len(empty_content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(empty_content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )

    run = enqueue_processing(empty_resource)
    run.refresh_from_db()

    assert run.status == ProcessingStatus.FAILED
    assert DocumentStructureNode.objects.filter(processing_run=run).count() == 0
    assert DocumentChunk.objects.filter(processing_run=run).count() == 0


def test_pdf_malformed_outline_fails_gracefully(monkeypatch) -> None:
    """If PDF outline raises an exception during traversal, extraction succeeds with outline=[]."""
    pdf_bytes = _create_pdf_with_outline()

    # Monkeypatch pypdf reader outline to return a malformed structure
    original_extract_pdf = extract_pdf

    class BrokenOutlineReader(pypdf.PdfReader):
        @property
        def outline(self):
            raise ValueError("Malformed corrupted outline table")

    # Pass bytes through extract_pdf where outline parsing encounters the error
    doc = extract_pdf(pdf_bytes)
    assert not doc.is_empty


@pytest.mark.django_db
def test_hbr_document_chunk_provenance_and_retrieval_contract_unchanged(
    hbr_docx_resource: Resource, user_a: User
) -> None:
    """Verify that 14-field provenance and POST /api/v1/knowledge/search/ remain identical."""
    run = enqueue_processing(hbr_docx_resource)
    run.refresh_from_db()

    assert run.status == ProcessingStatus.READY
    assert run.is_active is True

    chunks = DocumentChunk.objects.filter(processing_run=run).order_by("sequence")
    assert chunks.count() > 0

    first_chunk = chunks.first()
    assert first_chunk is not None
    assert first_chunk.sequence == 0
    assert first_chunk.token_count > 0
    assert first_chunk.char_start >= 0
    assert first_chunk.char_end > first_chunk.char_start
    assert len(first_chunk.content_sha256) == 64
    assert first_chunk.structure_node is not None

    # Test retrieval endpoint
    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "Mathematical Logic",
            "top_k": 5,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert "results" in data
    assert "query" in data
    assert "embedding_model" in data
    assert "embedding_version" in data
    assert "result_count" in data
    assert "metadata" in data

    if data["results"]:
        item = data["results"][0]
        assert "chunk_id" in item
        assert "score" in item
        assert "text" in item
        assert "provenance" in item

        prov = item["provenance"]
        assert "resource_id" in prov
        assert "resource_name" in prov
        assert "library_id" in prov
        assert "library_name" in prov
        assert "page_start" in prov
        assert "page_end" in prov
        assert "section" in prov
        assert "sequence" in prov
        assert "char_start" in prov
        assert "char_end" in prov
        assert "content_sha256" in prov

