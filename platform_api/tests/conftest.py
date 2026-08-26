"""Shared pytest fixtures for the Platform API test suite."""

from typing import TYPE_CHECKING

import pytest
from django.contrib.auth import get_user_model
from rest_framework.test import APIClient

from platform_api.apps.institutions.models import Institution, InstitutionStatus
from platform_api.apps.libraries.models import (
    Library,
    LibraryAccessPolicy,
    LibraryAccessRole,
    LibraryStatus,
    LibraryVisibility,
)
from platform_api.apps.memberships.models import (
    Membership,
    MembershipRole,
    MembershipStatus,
)
from platform_api.apps.resources.fake_storage import FakeStorage

if TYPE_CHECKING:
    from platform_api.apps.users.models import User


@pytest.fixture
def api_client() -> APIClient:
    """Return an unauthenticated DRF API client."""
    return APIClient()


@pytest.fixture
def user_a(db) -> "User":
    """Return User A."""
    user_model = get_user_model()
    return user_model.objects.create_user(
        email="user.a@example.com", password="password-a-123"
    )


@pytest.fixture
def user_b(db) -> "User":
    """Return User B."""
    user_model = get_user_model()
    return user_model.objects.create_user(
        email="user.b@example.com", password="password-b-123"
    )


@pytest.fixture
def institution_a(db) -> Institution:
    """Return Institution A."""
    return Institution.objects.create(
        name="Institution A",
        slug="institution-a",
        status=InstitutionStatus.ACTIVE,
    )


@pytest.fixture
def institution_b(db) -> Institution:
    """Return Institution B."""
    return Institution.objects.create(
        name="Institution B",
        slug="institution-b",
        status=InstitutionStatus.ACTIVE,
    )


@pytest.fixture
def client_a(user_a: "User") -> APIClient:
    """Return an API client authenticated as User A."""
    client = APIClient()
    client.force_authenticate(user=user_a)
    return client


@pytest.fixture
def client_b(user_b: "User") -> APIClient:
    """Return an API client authenticated as User B."""
    client = APIClient()
    client.force_authenticate(user=user_b)
    return client


@pytest.fixture
def membership_a(user_a: "User", institution_a: Institution) -> Membership:
    """Return User A's active membership in Institution A."""
    return Membership.objects.create(
        user=user_a,
        institution=institution_a,
        role=MembershipRole.TEACHER,
        status=MembershipStatus.ACTIVE,
    )


@pytest.fixture
def membership_b(user_b: "User", institution_b: Institution) -> Membership:
    """Return User B's active membership in Institution B."""
    return Membership.objects.create(
        user=user_b,
        institution=institution_b,
        role=MembershipRole.STUDENT,
        status=MembershipStatus.ACTIVE,
    )


@pytest.fixture
def admin_membership_a(user_a: "User", institution_a: Institution) -> Membership:
    """Return User A's active administrator membership in Institution A."""
    return Membership.objects.create(
        user=user_a,
        institution=institution_a,
        role=MembershipRole.ADMINISTRATOR,
        status=MembershipStatus.ACTIVE,
    )


@pytest.fixture
def admin_client_a(user_a: "User", admin_membership_a: Membership) -> APIClient:
    """Return an API client authenticated as User A, who is an institution admin."""
    client = APIClient()
    client.force_authenticate(user=user_a)
    return client


@pytest.fixture
def library_a(institution_a: Institution) -> Library:
    """Return a restricted active library in Institution A."""
    return Library.objects.create(
        institution=institution_a,
        name="Library A",
        slug="library-a",
        description="A test library.",
        status=LibraryStatus.ACTIVE,
        visibility=LibraryVisibility.RESTRICTED,
    )


@pytest.fixture
def library_b(institution_b: Institution) -> Library:
    """Return a restricted active library in Institution B."""
    return Library.objects.create(
        institution=institution_b,
        name="Library B",
        slug="library-b",
        description="Another test library.",
        status=LibraryStatus.ACTIVE,
        visibility=LibraryVisibility.RESTRICTED,
    )


@pytest.fixture
def discoverable_library_a(institution_a: Institution) -> Library:
    """Return a discoverable active library in Institution A."""
    return Library.objects.create(
        institution=institution_a,
        name="Discoverable Library A",
        slug="discoverable-library-a",
        status=LibraryStatus.ACTIVE,
        visibility=LibraryVisibility.DISCOVERABLE,
    )


@pytest.fixture
def library_admin_policy_a(user_a: "User", library_a: Library) -> LibraryAccessPolicy:
    """Return a library administrator access policy for User A on Library A."""
    return LibraryAccessPolicy.objects.create(
        library=library_a,
        user=user_a,
        role=LibraryAccessRole.ADMINISTRATOR,
    )


@pytest.fixture
def library_teacher_policy_a(user_a: "User", library_a: Library) -> LibraryAccessPolicy:
    """Return a teacher access policy for User A on Library A."""
    return LibraryAccessPolicy.objects.create(
        library=library_a,
        user=user_a,
        role=LibraryAccessRole.TEACHER,
    )


@pytest.fixture
def library_student_policy_a(user_a: "User", library_a: Library) -> LibraryAccessPolicy:
    """Return a student access policy for User A on Library A."""
    return LibraryAccessPolicy.objects.create(
        library=library_a,
        user=user_a,
        role=LibraryAccessRole.STUDENT,
    )


@pytest.fixture(autouse=True)
def _test_settings(settings) -> None:
    """Configure test defaults for storage, embeddings, and celery."""
    settings.OBJECT_STORAGE_BACKEND = (
        "platform_api.apps.resources.fake_storage.FakeStorage"
    )
    settings.EMBEDDING_PROVIDER_BACKEND = (
        "platform_api.apps.processing.embedding.fake_provider.FakeEmbeddingProvider"
    )
    settings.CELERY_TASK_ALWAYS_EAGER = True
    settings.CELERY_TASK_EAGER_PROPAGATES = True
    # The test client talks over plain http; relax the Secure flag so cookies
    # round-trip. Production defaults remain Secure=True.
    settings.REFRESH_COOKIE_SECURE = False
    settings.CSRF_COOKIE_SECURE = False
    FakeStorage.clear()


@pytest.fixture
def txt_bytes() -> bytes:
    """Return sample plain-text content."""
    return (
        b"Introduction to Biology\n\n"
        b"Biology is the study of living organisms and their "
        b"interactions with the environment.\n\n"
        b"Cell Structure\n\n"
        b"Cells are the basic structural and functional units of "
        b"all living organisms."
    )


@pytest.fixture
def pdf_bytes() -> bytes:
    """Return sample PDF content."""
    from io import BytesIO

    import pypdf

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    stream = BytesIO()
    writer.write(stream)
    return stream.getvalue()


@pytest.fixture
def docx_bytes() -> bytes:
    """Return sample DOCX content with headings and paragraphs."""
    from io import BytesIO

    import docx

    doc = docx.Document()
    doc.add_heading("Chapter 1: Quantum Physics", level=1)
    doc.add_paragraph(
        "Quantum physics is the study of matter and energy at the "
        "most fundamental level."
    )
    doc.add_paragraph("A central concept is wave-particle duality.")
    doc.add_heading("Chapter 2: Thermodynamics", level=1)
    doc.add_paragraph("Thermodynamics deals with heat, work, and temperature.")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()
