"""Adversarial and functional test suite for Phase 1 Hierarchical Structural Retrieval."""

import uuid
from io import BytesIO

import docx
import pytest
from rest_framework.test import APIClient

from platform_api.apps.knowledge.authentication import mint_delegated_token
from platform_api.apps.libraries.models import (
    Library,
    LibraryAccessPolicy,
)
from platform_api.apps.memberships.models import Membership
from platform_api.apps.processing.models import (
    DocumentChunk,
    DocumentStructureNode,
    ProcessingStatus,
)
from platform_api.apps.processing.services import enqueue_processing
from platform_api.apps.resources.checksum import sha256_checksum
from platform_api.apps.resources.models import Resource, ResourceStatus, ResourceType
from platform_api.apps.resources.object_key import generate_resource_object_key
from platform_api.apps.resources.storage import get_object_storage
from platform_api.apps.users.models import User


def _create_chemistry_textbook_docx() -> bytes:
    """Create a structured textbook containing both Chemical Kinetics and Thermodynamics chapters."""
    doc = docx.Document()

    # Chapter 1: Chemical Kinetics
    doc.add_heading("Chapter 1: Chemical Kinetics", level=1)
    doc.add_paragraph(
        "Chemical kinetics is the branch of chemistry that concerns the rates of chemical reactions, "
        "the factors that influence these rates, and the microscopic mechanisms by which reactions occur. "
        "Understanding reaction rates is essential for controlling industrial synthesis and biological processes."
    )
    doc.add_heading("Section 1.1: Reaction Rates and Temperature", level=2)
    doc.add_paragraph(
        "The rate of a chemical reaction increases significantly as temperature rises. "
        "Higher temperature increases the average kinetic energy and collision frequency of reactant molecules, "
        "allowing a greater fraction of molecules to possess energy greater than the activation energy barrier. "
        "For many homogeneous reactions, the reaction rate approximately doubles for every ten-degree Celsius rise in temperature."
    )
    doc.add_paragraph(
        "According to transition state theory and collision theory, the rate constant k varies exponentially with temperature "
        "following the Arrhenius equation: k = A * exp(-Ea / RT), where A is the pre-exponential frequency factor, "
        "Ea is the activation energy, R is the universal gas constant, and T is absolute temperature."
    )
    doc.add_paragraph(
        "Catalysts accelerate reaction rates by providing an alternative reaction pathway with a lower activation energy, "
        "without being consumed in the overall reaction."
    )

    # Chapter 8: Thermodynamics
    doc.add_heading("Chapter 8: Thermodynamics", level=1)
    doc.add_paragraph(
        "Thermodynamics deals with heat, work, entropy, and energy transformations in physical systems. "
        "It provides macroscopic relationships between state variables such as internal energy, enthalpy, and temperature."
    )
    doc.add_heading("Section 8.1: Temperature and Thermal Equilibrium", level=2)
    doc.add_paragraph(
        "Temperature is an intensive thermodynamic state variable reflecting the average kinetic energy of constituent particles. "
        "The Zeroth Law of Thermodynamics establishes that if two thermodynamic systems are each in thermal equilibrium with a third, "
        "they are in thermal equilibrium with each other."
    )
    doc.add_paragraph(
        "When two bodies at different initial temperatures are placed in diathermal contact, spontaneous heat flows from the hotter "
        "body to the colder body until thermal equilibrium is reached at a uniform final temperature, maximizing total system entropy."
    )

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _create_equilibrium_textbook_docx() -> bytes:
    """Create a structured textbook containing Chemical Equilibrium and Ideal Gas Laws chapters."""
    doc = docx.Document()

    # Chapter 3: Chemical Equilibrium
    doc.add_heading("Chapter 3: Chemical Equilibrium", level=1)
    doc.add_paragraph(
        "Dynamic equilibrium occurs in reversible chemical reactions when the forward reaction rate equals the reverse reaction rate. "
        "At dynamic chemical equilibrium, the concentrations of reactants and products remain constant over time."
    )
    doc.add_heading("Section 3.1: Le Chatelier's Principle and Pressure", level=2)
    doc.add_paragraph(
        "When pressure is increased in a gaseous chemical equilibrium system, the position of "
        "equilibrium shifts toward the side with the smaller number of moles of gas to relieve the imposed stress. "
        "Conversely, decreasing pressure shifts the equilibrium toward the side with more gaseous moles."
    )

    # Chapter 5: Ideal Gas Laws
    doc.add_heading("Chapter 5: Ideal Gas Laws", level=1)
    doc.add_paragraph(
        "The ideal gas equation PV = nRT describes the state of a hypothetical ideal gas under varying macroscopic conditions."
    )
    doc.add_heading("Section 5.1: Boyle's Law", level=2)
    doc.add_paragraph(
        "Boyle's law states that at constant temperature, the volume of a given mass of dry gas "
        "is inversely proportional to its absolute pressure. Increasing external pressure compresses the gas into a smaller volume."
    )

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


@pytest.fixture
def chemistry_textbook_resource(db, library_a: Library, user_a: User) -> Resource:
    """Upload and create a chemistry textbook resource."""
    content = _create_chemistry_textbook_docx()
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(
        key,
        BytesIO(content),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        len(content),
    )

    return Resource.objects.create(
        id=res_id,
        library=library_a,
        name="general_chemistry.docx",
        resource_type=ResourceType.DOCX,
        original_filename="general_chemistry.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=len(content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )


@pytest.fixture
def equilibrium_textbook_resource(db, library_a: Library, user_a: User) -> Resource:
    """Upload and create an equilibrium/gas textbook resource."""
    content = _create_equilibrium_textbook_docx()
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(
        key,
        BytesIO(content),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        len(content),
    )

    return Resource.objects.create(
        id=res_id,
        library=library_a,
        name="physical_chemistry.docx",
        resource_type=ResourceType.DOCX,
        original_filename="physical_chemistry.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=len(content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )


@pytest.mark.django_db
def test_adversarial_kinetics_vs_thermodynamics(
    chemistry_textbook_resource: Resource,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Querying temperature effect on reaction rate MUST retrieve Chemical Kinetics, NOT Thermodynamics."""
    run = enqueue_processing(chemistry_textbook_resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    # Adversarial Query: contains "temperature" heavily present in Thermodynamics chapter,
    # but the core concept is "rate of a chemical reaction".
    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "How does temperature affect the rate of a chemical reaction?",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    assert data["metadata"]["retrieval_strategy"] == "structural_section"
    assert data["metadata"]["candidate_sections_matched"] > 0

    # Verify that top result originates from Chemical Kinetics and NOT Thermodynamics
    top_section = data["results"][0]["provenance"]["section"] or ""
    assert "Chemical Kinetics" in top_section or "Reaction Rates" in top_section
    assert "Thermodynamics" not in top_section


@pytest.mark.django_db
def test_adversarial_equilibrium_vs_gas_laws(
    equilibrium_textbook_resource: Resource,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Querying pressure effect on equilibrium MUST retrieve Chemical Equilibrium, NOT Boyle's Law."""
    run = enqueue_processing(equilibrium_textbook_resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "Why does increasing pressure affect chemical equilibrium?",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    assert data["metadata"]["retrieval_strategy"] == "structural_section"

    top_section = data["results"][0]["provenance"]["section"] or ""
    assert "Equilibrium" in top_section or "Le Chatelier" in top_section
    assert "Boyle's Law" not in top_section



def _create_biology_textbook_docx() -> bytes:
    """Create a structured textbook containing Photosynthesis and Cellular Respiration chapters."""
    doc = docx.Document()

    # Chapter 4: Photosynthesis
    doc.add_heading("Chapter 4: Photosynthesis", level=1)
    doc.add_paragraph(
        "Photosynthesis is the process by which green plants transform light energy into chemical energy."
    )
    doc.add_heading("Section 4.1: Light-Dependent Reactions and Oxygen Evolution", level=2)
    doc.add_paragraph(
        "During the light-dependent reactions in the thylakoid membranes, photolysis of water occurs. "
        "Water molecules are split into hydrogen ions, electrons, and molecular oxygen (O2), which is "
        "released into the atmosphere as a byproduct."
    )

    # Chapter 7: Cellular Respiration
    doc.add_heading("Chapter 7: Cellular Respiration", level=1)
    doc.add_paragraph(
        "Cellular respiration is a metabolic pathway that breaks down glucose to produce ATP."
    )
    doc.add_heading("Section 7.1: Aerobic Respiration and Oxygen Consumption", level=2)
    doc.add_paragraph(
        "In aerobic respiration, oxygen is consumed as the terminal electron acceptor in the electron "
        "transport chain located in the inner mitochondrial membrane, producing water and ATP."
    )

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


@pytest.fixture
def biology_textbook_resource(db, library_a: Library, user_a: User) -> Resource:
    """Upload and create a biology textbook resource."""
    content = _create_biology_textbook_docx()
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(
        key,
        BytesIO(content),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        len(content),
    )

    return Resource.objects.create(
        id=res_id,
        library=library_a,
        name="general_biology.docx",
        resource_type=ResourceType.DOCX,
        original_filename="general_biology.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=len(content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )


@pytest.mark.django_db
def test_unstructured_document_vector_fallback(
    db,
    library_a: Library,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Plain text document without structure nodes gracefully falls back to global hybrid search."""
    plain_content = (
        b"Python is an interpreted, high-level programming language.\n\n"
        b"Variables are created when you assign a value to them."
    )
    res_id = uuid.uuid4()
    key = generate_resource_object_key(library_a.pk, res_id)
    storage = get_object_storage()
    storage.upload(key, BytesIO(plain_content), "text/plain", len(plain_content))

    resource = Resource.objects.create(
        id=res_id,
        library=library_a,
        name="python_intro.txt",
        resource_type=ResourceType.TXT,
        original_filename="python_intro.txt",
        content_type="text/plain",
        size=len(plain_content),
        object_key=key,
        checksum=sha256_checksum(BytesIO(plain_content)),
        status=ResourceStatus.READY,
        created_by=user_a,
    )

    run = enqueue_processing(resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "How are variables created in Python?",
            "top_k": 2,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    assert data["metadata"]["retrieval_strategy"] == "global_hybrid"


@pytest.mark.django_db
def test_adversarial_photosynthesis_vs_respiration(
    biology_textbook_resource: Resource,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Querying why plants release oxygen MUST retrieve Photosynthesis, NOT Cellular Respiration."""
    run = enqueue_processing(biology_textbook_resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "Why does a plant release oxygen during photosynthesis?",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    assert data["metadata"]["retrieval_strategy"] == "structural_section"

    top_section = data["results"][0]["provenance"]["section"] or ""
    assert "Photosynthesis" in top_section or "Light-Dependent" in top_section
    assert "Respiration" not in top_section


@pytest.mark.django_db
def test_lexical_semantic_complementarity_ranking(
    chemistry_textbook_resource: Resource,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Specific technical terminology (Arrhenius equation) boosts the exact formula chunk to rank #1."""
    run = enqueue_processing(chemistry_textbook_resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "Arrhenius equation rate constant temperature dependence",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0

    top_result = data["results"][0]
    # Exact keyword match for Arrhenius equation should be present in top text
    assert "Arrhenius" in top_result["text"]
    assert top_result["score"] > 0.25



@pytest.mark.django_db
def test_cross_library_structural_authorization_isolation(
    chemistry_textbook_resource: Resource, user_b: User
) -> None:
    """User B (who is not authorized on Library A) cannot target or retrieve Library A structure."""
    run = enqueue_processing(chemistry_textbook_resource)
    run.refresh_from_db()
    assert run.status == ProcessingStatus.READY

    client = APIClient()
    # User B has no access to Library A
    token = mint_delegated_token(user_id=user_b.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "How does temperature affect the rate of a chemical reaction?",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] == 0
    assert len(data["results"]) == 0


# ==============================================================================
# STAGE 4 — BOUNDED CONTEXT EXPANSION TESTS
# ==============================================================================

from platform_api.apps.knowledge.context_expansion import expand_retrieval_context
from platform_api.apps.knowledge.dto import ProvenanceDTO, SearchResultItemDTO
from platform_api.apps.knowledge.policies import EffectiveRetrievalScope
from platform_api.apps.processing.models import ProcessingRun


def _create_test_chunk(
    run: ProcessingRun,
    library: Library,
    resource: Resource,
    seq: int,
    text: str,
    structure_node: DocumentStructureNode | None = None,
    section: str | None = None,
) -> DocumentChunk:
    return DocumentChunk.objects.create(
        processing_run=run,
        library=library,
        resource=resource,
        sequence=seq,
        text=text,
        token_count=len(text.split()),
        char_start=seq * 100,
        char_end=(seq + 1) * 100,
        page_start=1,
        page_end=1,
        structure_node=structure_node,
        section=section or (structure_node.title if structure_node else None),
        content_sha256=f"hash-seq-{seq}",
    )


@pytest.mark.django_db
def test_context_expansion_previous_and_next(
    db, library_a: Library, user_a: User
) -> None:
    """Core chunk 11 expands to include adjacent sequence 10 and sequence 12 in order [10, 11, 12]."""
    res = Resource.objects.create(
        library=library_a,
        name="test_doc.pdf",
        resource_type=ResourceType.PDF,
        original_filename="test_doc.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/test_doc.pdf",
        checksum="hash-res-1",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    node = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Chapter 1: Intro",
        normalized_title="chapter 1: intro",
        level=1,
        sequence=0,
    )

    c10 = _create_test_chunk(run, library_a, res, 10, "Paragraph 10 text", node)
    c11 = _create_test_chunk(run, library_a, res, 11, "Paragraph 11 text (core)", node)
    c12 = _create_test_chunk(run, library_a, res, 12, "Paragraph 12 text", node)

    core_dto = SearchResultItemDTO(
        chunk_id=c11.id,
        score=0.95,
        text=c11.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=node.title,
            sequence=11,
            char_start=1100,
            char_end=1200,
            content_sha256="hash-11",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 3
    assert {x.provenance.sequence for x in expanded} == {10, 11, 12}
    assert expanded[0].chunk_id == c11.id
    assert expanded[0].score == 0.95


@pytest.mark.django_db
def test_context_expansion_beginning_of_section(
    db, library_a: Library, user_a: User
) -> None:
    """Core chunk 20 (first chunk of Section 2) expands only to 21, rejecting chunk 19 from Section 1."""
    res = Resource.objects.create(
        library=library_a,
        name="test_book.pdf",
        resource_type=ResourceType.PDF,
        original_filename="test_book.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/test_book.pdf",
        checksum="hash-res-2",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    node1 = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Section 1",
        normalized_title="section 1",
        level=2,
        sequence=0,
    )
    node2 = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Section 2",
        normalized_title="section 2",
        level=2,
        sequence=1,
    )

    c19 = _create_test_chunk(run, library_a, res, 19, "Section 1 end", node1)
    c20 = _create_test_chunk(run, library_a, res, 20, "Section 2 start (core)", node2)
    c21 = _create_test_chunk(run, library_a, res, 21, "Section 2 middle", node2)

    core_dto = SearchResultItemDTO(
        chunk_id=c20.id,
        score=0.92,
        text=c20.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=node2.title,
            sequence=20,
            char_start=2000,
            char_end=2100,
            content_sha256="hash-20",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 2
    assert {x.provenance.sequence for x in expanded} == {20, 21}
    assert expanded[0].chunk_id == c20.id
    assert not any(x.provenance.sequence == 19 for x in expanded)


@pytest.mark.django_db
def test_context_expansion_end_of_section(
    db, library_a: Library, user_a: User
) -> None:
    """Core chunk 31 (last chunk of Section 1) expands only to 30, rejecting chunk 32 from Section 2."""
    res = Resource.objects.create(
        library=library_a,
        name="test_book3.pdf",
        resource_type=ResourceType.PDF,
        original_filename="test_book3.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/test_book3.pdf",
        checksum="hash-res-3",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    node1 = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Section 1",
        normalized_title="section 1",
        level=2,
        sequence=0,
    )
    node2 = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Section 2",
        normalized_title="section 2",
        level=2,
        sequence=1,
    )

    c30 = _create_test_chunk(run, library_a, res, 30, "Section 1 middle", node1)
    c31 = _create_test_chunk(run, library_a, res, 31, "Section 1 end (core)", node1)
    c32 = _create_test_chunk(run, library_a, res, 32, "Section 2 start", node2)

    core_dto = SearchResultItemDTO(
        chunk_id=c31.id,
        score=0.90,
        text=c31.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=node1.title,
            sequence=31,
            char_start=3100,
            char_end=3200,
            content_sha256="hash-31",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 2
    assert {x.provenance.sequence for x in expanded} == {30, 31}
    assert expanded[0].chunk_id == c31.id
    assert not any(x.provenance.sequence == 32 for x in expanded)


@pytest.mark.django_db
def test_context_expansion_unstructured_document(
    db, library_a: Library, user_a: User
) -> None:
    """Unstructured document expands sequence ± 1 within the same resource [50, 51, 52]."""
    res = Resource.objects.create(
        library=library_a,
        name="notes.txt",
        resource_type=ResourceType.TXT,
        original_filename="notes.txt",
        content_type="text/plain",
        size=1024,
        object_key="lib_a/notes.txt",
        checksum="hash-res-unstruct",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    c50 = _create_test_chunk(run, library_a, res, 50, "Unstructured 50", None)
    c51 = _create_test_chunk(run, library_a, res, 51, "Unstructured 51 (core)", None)
    c52 = _create_test_chunk(run, library_a, res, 52, "Unstructured 52", None)

    core_dto = SearchResultItemDTO(
        chunk_id=c51.id,
        score=0.88,
        text=c51.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=None,
            sequence=51,
            char_start=5100,
            char_end=5200,
            content_sha256="hash-51",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 3
    assert {x.provenance.sequence for x in expanded} == {50, 51, 52}
    assert expanded[0].chunk_id == c51.id


@pytest.mark.django_db
def test_context_expansion_deduplication(
    db, library_a: Library, user_a: User
) -> None:
    """Overlapping core results 51 and 52 deduplicate cleanly into [50, 51, 52, 53]."""
    res = Resource.objects.create(
        library=library_a,
        name="dedup.pdf",
        resource_type=ResourceType.PDF,
        original_filename="dedup.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/dedup.pdf",
        checksum="hash-res-dedup",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    c50 = _create_test_chunk(run, library_a, res, 50, "Chunk 50", None)
    c51 = _create_test_chunk(run, library_a, res, 51, "Chunk 51 (core 1)", None)
    c52 = _create_test_chunk(run, library_a, res, 52, "Chunk 52 (core 2)", None)
    c53 = _create_test_chunk(run, library_a, res, 53, "Chunk 53", None)

    core1 = SearchResultItemDTO(
        chunk_id=c51.id,
        score=0.95,
        text=c51.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=None,
            sequence=51,
            char_start=5100,
            char_end=5200,
            content_sha256="hash-51",
        ),
    )
    core2 = SearchResultItemDTO(
        chunk_id=c52.id,
        score=0.90,
        text=c52.text,
        provenance=ProvenanceDTO(
            resource_id=res.id,
            resource_name=res.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=None,
            sequence=52,
            char_start=5200,
            char_end=5300,
            content_sha256="hash-52",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core1, core2], scope, context_window=1)

    assert len(expanded) == 4
    assert {x.provenance.sequence for x in expanded} == {50, 51, 52, 53}
    # Verify no duplicate chunk IDs
    ids = [x.chunk_id for x in expanded]
    assert len(ids) == len(set(ids))



@pytest.mark.django_db
def test_context_expansion_cross_resource_isolation(
    db, library_a: Library, user_a: User
) -> None:
    """Resource A sequence 10 cannot expand into Resource B sequence 9."""
    res_a = Resource.objects.create(
        library=library_a,
        name="doc_a.pdf",
        resource_type=ResourceType.PDF,
        original_filename="doc_a.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/doc_a.pdf",
        checksum="hash-res-a",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run_a = ProcessingRun.objects.create(
        resource=res_a,
        library=library_a,
        source_checksum=res_a.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    res_b = Resource.objects.create(
        library=library_a,
        name="doc_b.pdf",
        resource_type=ResourceType.PDF,
        original_filename="doc_b.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/doc_b.pdf",
        checksum="hash-res-b",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run_b = ProcessingRun.objects.create(
        resource=res_b,
        library=library_a,
        source_checksum=res_b.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    ca10 = _create_test_chunk(run_a, library_a, res_a, 10, "Doc A chunk 10 (core)", None)
    cb9 = _create_test_chunk(run_b, library_a, res_b, 9, "Doc B chunk 9 (unrelated)", None)

    core_dto = SearchResultItemDTO(
        chunk_id=ca10.id,
        score=0.90,
        text=ca10.text,
        provenance=ProvenanceDTO(
            resource_id=res_a.id,
            resource_name=res_a.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=None,
            sequence=10,
            char_start=1000,
            char_end=1100,
            content_sha256="hash-10",
        ),
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 1
    assert expanded[0].chunk_id == ca10.id
    assert not any(x.chunk_id == cb9.id for x in expanded)


@pytest.mark.django_db
def test_context_expansion_cross_library_isolation(
    db, library_a: Library, library_b: Library, user_a: User, user_b: User
) -> None:
    """Library A chunk 100 cannot expand into Library B chunk 99."""
    res_a = Resource.objects.create(
        library=library_a,
        name="lib_a_doc.pdf",
        resource_type=ResourceType.PDF,
        original_filename="lib_a_doc.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/lib_a_doc.pdf",
        checksum="hash-lib-a",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run_a = ProcessingRun.objects.create(
        resource=res_a,
        library=library_a,
        source_checksum=res_a.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    res_b = Resource.objects.create(
        library=library_b,
        name="lib_b_doc.pdf",
        resource_type=ResourceType.PDF,
        original_filename="lib_b_doc.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_b/lib_b_doc.pdf",
        checksum="hash-lib-b",
        status=ResourceStatus.READY,
        created_by=user_b,
    )
    run_b = ProcessingRun.objects.create(
        resource=res_b,
        library=library_b,
        source_checksum=res_b.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    ca100 = _create_test_chunk(run_a, library_a, res_a, 100, "Lib A chunk 100", None)
    cb99 = _create_test_chunk(run_b, library_b, res_b, 99, "Lib B chunk 99", None)

    core_dto = SearchResultItemDTO(
        chunk_id=ca100.id,
        score=0.90,
        text=ca100.text,
        provenance=ProvenanceDTO(
            resource_id=res_a.id,
            resource_name=res_a.name,
            library_id=library_a.id,
            library_name=library_a.name,
            page_start=1,
            page_end=1,
            section=None,
            sequence=100,
            char_start=10000,
            char_end=10100,
            content_sha256="hash-100",
        ),
    )

    # Scope only authorizes Library A
    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    expanded = expand_retrieval_context([core_dto], scope, context_window=1)

    assert len(expanded) == 1
    assert expanded[0].chunk_id == ca100.id
    assert not any(x.provenance.library_id == library_b.id for x in expanded)


# ==============================================================================
# STAGE 5 — BACK-OF-BOOK INDEX EXTRACTION & GUIDED RETRIEVAL TESTS
# ==============================================================================

from platform_api.apps.knowledge.index_search import find_candidate_index_pages
from platform_api.apps.processing.extractors import ExtractedPage, OutlineNode
from platform_api.apps.processing.index_parser import (
    detect_index_pages,
    normalize_index_term,
    parse_index_entries,
    parse_page_references,
)
from platform_api.apps.processing.embedding import get_embedding_provider
from platform_api.apps.processing.models import BookIndexEntry, ChunkEmbedding




def test_index_term_normalization() -> None:
    """Term normalization handles case, whitespace, Unicode NFC, and punctuation consistently."""
    assert normalize_index_term("Chemical Kinetics") == "chemical kinetics"
    assert normalize_index_term("chemical kinetics") == "chemical kinetics"
    assert normalize_index_term("CHEMICAL   KINETICS") == "chemical kinetics"
    assert normalize_index_term("  activation energy... ") == "activation energy"
    assert normalize_index_term("reaction-rate") == "reaction rate"


def test_index_page_reference_parsing() -> None:
    """Page reference parsing correctly extracts comma lists, hyphen ranges, and dot leaders."""
    assert parse_page_references("6, 27, 103") == [6, 27, 103]
    assert parse_page_references("42–45") == [42, 43, 44, 45]
    assert parse_page_references("5, 7-9, 12") == [5, 7, 8, 9, 12]
    assert parse_page_references("..... 45, 47") == [45, 47]


def test_index_detection_and_rejection() -> None:
    """Index detector detects valid subject indexes and strictly rejects non-index sections."""
    index_page = ExtractedPage(
        page=100,
        text="SUBJECT INDEX\nactivation energy ........ 6, 27\nchemical equilibrium ..... 42, 51–53\nreaction rate ............ 4–9\ntemperature .............. 6, 81",
    )
    biblio_page = ExtractedPage(
        page=99,
        text="BIBLIOGRAPHY\n1. Smith, J. Chemical Kinetics. 2020.\n2. Doe, A. Thermodynamics. 2021.",
    )
    glossary_page = ExtractedPage(
        page=98,
        text="GLOSSARY\nCatalyst: A substance that increases rate of reaction.\nEnthalpy: Total heat content.",
    )

    detected = detect_index_pages([glossary_page, biblio_page, index_page])
    assert len(detected) == 1
    assert detected[0].page == 100


def test_index_entry_parsing_simple_ranges_and_subentries() -> None:
    """Index parser extracts parent terms, ranges, and indented subentries accurately."""
    page = ExtractedPage(
        page=100,
        text="""SUBJECT INDEX
activation energy, 6, 27, 103
chemical kinetics, 4–8
reaction rates
    factors affecting, 6–8
    temperature effects, 7
""",
    )
    entries = parse_index_entries([page])
    assert len(entries) == 4

    # Entry 1: activation energy
    assert entries[0].term == "activation energy"
    assert entries[0].target_physical_pages == [6, 27, 103]

    # Entry 2: chemical kinetics
    assert entries[1].term == "chemical kinetics"
    assert entries[1].target_physical_pages == [4, 5, 6, 7, 8]

    # Entry 3: subentry factors affecting
    assert entries[2].term == "reaction rates"
    assert entries[2].subterm == "factors affecting"
    assert entries[2].normalized_term == "reaction rates factors affecting"
    assert entries[2].target_physical_pages == [6, 7, 8]

    # Entry 4: subentry temperature effects
    assert entries[3].term == "reaction rates"
    assert entries[3].subterm == "temperature effects"
    assert entries[3].target_physical_pages == [7]


@pytest.mark.django_db
def test_index_search_multi_concept_intersection(
    db, library_a: Library, user_a: User
) -> None:
    """Query with multiple concepts ('temperature' + 'reaction rate') resolves to candidate intersection page 7."""
    res = Resource.objects.create(
        library=library_a,
        name="chemistry.pdf",
        resource_type=ResourceType.PDF,
        original_filename="chemistry.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/chemistry.pdf",
        checksum="hash-chem-idx",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    # Index entry 1: reaction rate -> pages 5..9
    BookIndexEntry.objects.create(
        processing_run=run,
        resource=res,
        term="reaction rate",
        normalized_term="reaction rate",
        raw_page_references="5–9",
        target_physical_pages=[5, 6, 7, 8, 9],
    )
    # Index entry 2: temperature -> pages 7, 81
    BookIndexEntry.objects.create(
        processing_run=run,
        resource=res,
        term="temperature",
        normalized_term="temperature",
        raw_page_references="7, 81",
        target_physical_pages=[7, 81],
    )

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)
    pages = find_candidate_index_pages("How does temperature affect reaction rate?", scope)

    # Intersection of [5..9] and [7, 81] is [7]!
    assert pages == [7]


@pytest.mark.django_db
def test_index_guided_adversarial_kinetics_vs_thermodynamics(
    db,
    library_a: Library,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """Index-guided retrieval directs search to kinetics page 7 over thermodynamics page 81."""
    res = Resource.objects.create(
        library=library_a,
        name="adv_chem.pdf",
        resource_type=ResourceType.PDF,
        original_filename="adv_chem.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/adv_chem.pdf",
        checksum="hash-adv-chem",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    provider = get_embedding_provider()
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        embedding_dimensions=provider.dimensions,
        status=ProcessingStatus.READY,
        is_active=True,
    )


    # TOC Section 1: Chemical Kinetics (pages 4-10)
    node_kinetics = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Chapter 1: Chemical Kinetics",
        normalized_title="chapter 1: chemical kinetics",
        level=1,
        page_start=4,
        page_end=10,
        sequence=0,
    )
    # TOC Section 2: Thermodynamics (pages 80-100)
    node_thermo = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Chapter 8: Thermodynamics",
        normalized_title="chapter 8: thermodynamics",
        level=1,
        page_start=80,
        page_end=100,
        sequence=1,
    )

    # Create chunk in kinetics on page 7
    c_kinetics = DocumentChunk.objects.create(
        processing_run=run,
        library=library_a,
        resource=res,
        sequence=1,
        text="The rate of reaction increases as temperature rises. Molecular collisions occur with higher frequency.",
        token_count=15,
        char_start=0,
        char_end=100,
        page_start=7,
        page_end=7,
        structure_node=node_kinetics,
        section="Chapter 1: Chemical Kinetics",
        content_sha256="hash-kin-7",
    )
    ChunkEmbedding.objects.create(
        chunk=c_kinetics,
        vector=[0.0] * provider.dimensions,
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        dimensions=provider.dimensions,
    )

    # Create chunk in thermo on page 81
    c_thermo = DocumentChunk.objects.create(
        processing_run=run,
        library=library_a,
        resource=res,
        sequence=2,
        text="Thermodynamic temperature dictates entropy changes and Gibbs free energy in closed systems.",
        token_count=13,
        char_start=100,
        char_end=200,
        page_start=81,
        page_end=81,
        structure_node=node_thermo,
        section="Chapter 8: Thermodynamics",
        content_sha256="hash-thm-81",
    )
    ChunkEmbedding.objects.create(
        chunk=c_thermo,
        vector=[0.0] * provider.dimensions,
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        dimensions=provider.dimensions,
    )


    # Index entry pointing to kinetics page 7
    BookIndexEntry.objects.create(
        processing_run=run,
        resource=res,
        term="reaction rate",
        normalized_term="reaction rate",
        raw_page_references="7",
        target_physical_pages=[7],
    )
    BookIndexEntry.objects.create(
        processing_run=run,
        resource=res,
        term="temperature",
        normalized_term="temperature",
        raw_page_references="7, 81",
        target_physical_pages=[7, 81],
    )

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "How does temperature affect the rate of a reaction?",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    assert data["metadata"]["retrieval_strategy"] == "index_guided_hybrid"
    # Kinetics chunk must rank #1
    top_result = data["results"][0]
    assert "Molecular collisions" in top_result["text"]
    assert top_result["provenance"]["page_start"] == 7


@pytest.mark.django_db
def test_index_guided_cross_library_isolation(
    db, library_a: Library, library_b: Library, user_a: User, user_b: User
) -> None:
    """User B (Institution B) cannot query or resolve index entries belonging to Library A."""
    res_a = Resource.objects.create(

        library=library_a,
        name="private_a.pdf",
        resource_type=ResourceType.PDF,
        original_filename="private_a.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/private_a.pdf",
        checksum="hash-priv-a",
        status=ResourceStatus.READY,
        created_by=user_a,
    )

    run_a = ProcessingRun.objects.create(
        resource=res_a,
        library=library_a,
        source_checksum=res_a.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )
    BookIndexEntry.objects.create(
        processing_run=run_a,
        resource=res_a,
        term="quantum entanglement",
        normalized_term="quantum entanglement",
        raw_page_references="42",
        target_physical_pages=[42],
    )

    # User B scope only has Library B
    scope_b = EffectiveRetrievalScope(frozenset([library_b.id]), None)
    pages = find_candidate_index_pages("quantum entanglement", scope_b)

    assert len(pages) == 0


@pytest.mark.django_db
def test_index_toc_multiple_references_intersection(
    db,
    library_a: Library,
    user_a: User,
    membership_a: Membership,
    library_student_policy_a: LibraryAccessPolicy,
) -> None:
    """TOC section (Chemical Kinetics 4-18) intersects multiple index references (6, 27, 103) to prioritize page 6."""
    res = Resource.objects.create(
        library=library_a,
        name="multi_ref.pdf",
        resource_type=ResourceType.PDF,
        original_filename="multi_ref.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/multi_ref.pdf",
        checksum="hash-multi-ref",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    provider = get_embedding_provider()
    run = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        embedding_dimensions=provider.dimensions,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    node_kinetics = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Chapter 1: Chemical Kinetics",
        normalized_title="chapter 1: chemical kinetics",
        level=1,
        page_start=4,
        page_end=18,
        sequence=0,
    )
    node_electro = DocumentStructureNode.objects.create(
        processing_run=run,
        resource=res,
        library=library_a,
        title="Chapter 3: Electrochemistry",
        normalized_title="chapter 3: electrochemistry",
        level=1,
        page_start=25,
        page_end=35,
        sequence=1,
    )

    c_kin = DocumentChunk.objects.create(
        processing_run=run,
        library=library_a,
        resource=res,
        sequence=1,
        text="Activation energy represents the energy barrier that reactants must overcome to form products in chemical kinetics.",
        token_count=16,
        char_start=0,
        char_end=100,
        page_start=6,
        page_end=6,
        structure_node=node_kinetics,
        section="Chapter 1: Chemical Kinetics",
        content_sha256="hash-kin-6",
    )
    ChunkEmbedding.objects.create(
        chunk=c_kin,
        vector=[0.0] * provider.dimensions,
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        dimensions=provider.dimensions,
    )

    c_elec = DocumentChunk.objects.create(
        processing_run=run,
        library=library_a,
        resource=res,
        sequence=2,
        text="Electrochemical overpotential and activation energy in galvanic cell electrodes.",
        token_count=10,
        char_start=100,
        char_end=200,
        page_start=27,
        page_end=27,
        structure_node=node_electro,
        section="Chapter 3: Electrochemistry",
        content_sha256="hash-elec-27",
    )
    ChunkEmbedding.objects.create(
        chunk=c_elec,
        vector=[0.0] * provider.dimensions,
        embedding_model=provider.model_id,
        embedding_version=provider.embedding_version,
        dimensions=provider.dimensions,
    )

    BookIndexEntry.objects.create(
        processing_run=run,
        resource=res,
        term="activation energy",
        normalized_term="activation energy",
        raw_page_references="6, 27, 103",
        target_physical_pages=[6, 27, 103],
    )

    client = APIClient()
    token = mint_delegated_token(user_id=user_a.pk)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {token}")

    response = client.post(
        "/api/v1/knowledge/search/",
        {
            "query": "Chemical kinetics activation energy barrier",
            "top_k": 3,
        },
        format="json",
    )

    assert response.status_code == 200
    data = response.json()
    assert data["result_count"] > 0
    top = data["results"][0]
    assert top["provenance"]["page_start"] == 6
    assert "energy barrier" in top["text"]


@pytest.mark.django_db
def test_reprocessing_purges_old_index_entries(
    db, library_a: Library, user_a: User
) -> None:
    """When a resource is reprocessed, old index entries are replaced and only active run participates in search."""
    from platform_api.apps.processing.indexing import write_chunks_and_embeddings, activate_run

    res = Resource.objects.create(
        library=library_a,
        name="reprocess.pdf",
        resource_type=ResourceType.PDF,
        original_filename="reprocess.pdf",
        content_type="application/pdf",
        size=1024,
        object_key="lib_a/reprocess.pdf",
        checksum="hash-reproc-1",
        status=ResourceStatus.READY,
        created_by=user_a,
    )
    run1 = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum=res.checksum,
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.READY,
        is_active=True,
    )

    # Old run had entry 'thermodynamic entropy'
    BookIndexEntry.objects.create(
        processing_run=run1,
        resource=res,
        term="thermodynamic entropy",
        normalized_term="thermodynamic entropy",
        raw_page_references="10",
        target_physical_pages=[10],
    )

    # Create Run 2 (reprocessing)
    run2 = ProcessingRun.objects.create(
        resource=res,
        library=library_a,
        source_checksum="hash-reproc-2",
        pipeline_version="1",
        extractor_version="1",
        chunker_version="1",
        embedding_model="fake-model",
        embedding_version="1",
        embedding_dimensions=1536,
        status=ProcessingStatus.PROCESSING,
        is_active=False,
    )

    new_page = ExtractedPage(
        page=50,
        text="SUBJECT INDEX\nquantum mechanics ........ 25\n",
    )

    write_chunks_and_embeddings(
        run=run2,
        chunks=[],
        vectors=[],
        extracted_pages=[new_page],
    )
    activate_run(run2)

    scope = EffectiveRetrievalScope(frozenset([library_a.id]), None)

    # Old concept should no longer be returned because run1 is deactivated
    old_pages = find_candidate_index_pages("thermodynamic entropy", scope)
    assert len(old_pages) == 0

    # New concept should be returned from run2
    new_pages = find_candidate_index_pages("quantum mechanics", scope)
    assert new_pages == [25]






