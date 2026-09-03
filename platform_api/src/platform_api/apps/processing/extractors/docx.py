"""DOCX extractor using python-docx."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

import docx

from . import (
    ExtractedDocument,
    ExtractedPage,
    ExtractionError,
    OutlineNode,
    StructureNodeType,
)

EXTRACTOR_VERSION_DOCX = "docx-hbr1"


def _get_heading_level(style_name: str) -> int | None:
    """Extract 1-based hierarchy level from DOCX heading style name."""
    s = style_name.strip()
    lower = s.lower()
    if lower == "title":
        return 1
    if lower == "subtitle":
        return 2
    m = re.match(r"^heading\s*(\d+)", s, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return None


def _classify_docx_node_type(title: str, level: int) -> str:
    """Classify structural node type for DOCX heading."""
    t = title.strip().lower()
    if re.match(r"^(?:part\b|volume\b|book\b)", t):
        return StructureNodeType.PART
    if re.match(r"^(?:chapter\b|ch\.\s*\d+)", t):
        return StructureNodeType.CHAPTER
    if re.match(r"^(?:section\b|sec\.\s*\d+)", t):
        return StructureNodeType.SECTION
    if re.match(r"^(?:subsection\b)", t):
        return StructureNodeType.SUBSECTION
    if re.match(r"^(?:appendix\b)", t):
        return StructureNodeType.APPENDIX
    if re.match(
        r"^(?:preface\b|foreword\b|introduction\b|prologue\b|contents\b|table of contents\b|acknowledgements?\b)",
        t,
    ):
        return StructureNodeType.FRONT_MATTER
    if re.match(
        r"^(?:epilogue\b|afterword\b|glossary\b|bibliography\b|references\b|index\b)",
        t,
    ):
        return StructureNodeType.BACK_MATTER
    if level == 1:
        return (
            StructureNodeType.CHAPTER
            if "chapter" in t
            else StructureNodeType.SECTION
        )
    if level == 2:
        return StructureNodeType.SECTION
    if level >= 3:
        return StructureNodeType.SUBSECTION
    return StructureNodeType.OTHER



def extract_docx(content: bytes) -> ExtractedDocument:
    """Extract paragraphs and structural headings from DOCX bytes.

    Args:
        content: Raw DOCX bytes.

    Returns:
        ExtractedDocument containing paragraphs with heading associations and outline.

    Raises:
        ExtractionError: If the DOCX archive is invalid or unreadable.
    """
    try:
        stream = BytesIO(content)
        doc = docx.Document(stream)

        pages: list[ExtractedPage] = []
        current_heading: str | None = None

        root_nodes: list[dict[str, Any]] = []
        stack: list[dict[str, Any]] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = getattr(paragraph.style, "name", "") or ""
            level = _get_heading_level(style_name)

            if level is not None:
                current_heading = text
                node_type = _classify_docx_node_type(text, level)
                node_data: dict[str, Any] = {
                    "title": text,
                    "level": level,
                    "node_type": node_type,
                    "source": "heading_style",
                    "children": [],
                }
                while stack and stack[-1]["level"] >= level:
                    stack.pop()

                if stack:
                    stack[-1]["children"].append(node_data)
                else:
                    root_nodes.append(node_data)
                stack.append(node_data)

                pages.append(
                    ExtractedPage(
                        page=None,
                        text=text,
                        heading=current_heading,
                    )
                )
            else:
                pages.append(
                    ExtractedPage(
                        page=None,
                        text=text,
                        heading=current_heading,
                    )
                )

        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                row_texts = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_texts:
                    pages.append(
                        ExtractedPage(
                            page=None,
                            text=" | ".join(row_texts),
                            heading=current_heading,
                        )
                    )

        def _build_outline_nodes(
            nodes_data: list[dict[str, Any]], seq_counter: list[int]
        ) -> list[OutlineNode]:
            result: list[OutlineNode] = []
            for nd in nodes_data:
                seq = seq_counter[0]
                seq_counter[0] += 1
                children = (
                    _build_outline_nodes(nd["children"], seq_counter)
                    if nd["children"]
                    else []
                )
                result.append(
                    OutlineNode(
                        title=nd["title"],
                        level=nd["level"],
                        node_type=nd["node_type"],
                        page_start=None,
                        page_end=None,
                        sequence=seq,
                        source=nd["source"],
                        confidence=1.0,
                        metadata={},
                        children=children,
                    )
                )
            return result

        seq_counter = [0]
        outline = _build_outline_nodes(root_nodes, seq_counter)
        page_labels = [
            (p.page, str(p.page), "default")
            for p in pages
            if p.page is not None
        ]

        return ExtractedDocument(
            pages=pages,
            outline=outline,
            page_labels=page_labels,
        )
    except Exception as exc:

        raise ExtractionError(f"Failed to extract text from DOCX: {exc}") from exc

